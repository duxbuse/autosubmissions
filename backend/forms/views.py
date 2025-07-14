from rest_framework import viewsets
from rest_framework.response import Response
from rest_framework.decorators import action
from django.shortcuts import get_object_or_404
from .models import Form, Question, Submission, Answer
from .serializers import FormSerializer, SubmissionSerializer
import json
from django.conf import settings
from pathlib import Path
import logging
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

class FormViewSet(viewsets.ModelViewSet):
    queryset = Form.objects.all()
    serializer_class = FormSerializer

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance)
        return Response(serializer.data)

    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        data = request.data
        logger = logging.getLogger(__name__)
        logger.info(f"[FormViewSet.update] Incoming payload: {json.dumps(data, indent=2)}")
        # Save form and questions/options
        try:
            serializer = self.get_serializer(instance, data=data, partial=True)
            serializer.is_valid(raise_exception=True)
        except Exception as e:
            logger.error(f"Serializer error in FormViewSet.update: {e}\nData: {data}\nErrors: {getattr(e, 'detail', str(e))}")
            return Response({'detail': 'Invalid data', 'errors': getattr(e, 'detail', str(e))}, status=400)
        self.perform_update(serializer)
        # Log serialized data after update
        logger.info(f"[FormViewSet.update] Serialized response: {json.dumps(serializer.data, indent=2)}")
        # Compare payload and serialized data for divergence
        if json.dumps(data, sort_keys=True) != json.dumps(serializer.data, sort_keys=True):
            logger.warning(f"[FormViewSet.update] Payload and serialized data diverge!\nPayload: {json.dumps(data, indent=2)}\nSerialized: {json.dumps(serializer.data, indent=2)}")
        # Write to JSON file
        forms_dir = Path(settings.BASE_DIR) / 'form_json'
        forms_dir.mkdir(exist_ok=True)
        file_path = forms_dir / f'form_{instance.id}.json'
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(serializer.data, f, ensure_ascii=False, indent=2)
        return Response(serializer.data)


class SubmissionViewSet(viewsets.ModelViewSet):
    queryset = Submission.objects.all()
    serializer_class = SubmissionSerializer

    def destroy(self, request, *args, **kwargs):
        logger = logging.getLogger(__name__)
        submission = self.get_object()
        logger.info(f"[SubmissionViewSet.destroy] Deleting submission id={submission.pk}, client_name={submission.client_name}")
        # Optionally, dump the submission to a JSON file before deletion for audit/debug
        try:
            from django.conf import settings
            from pathlib import Path
            import datetime
            submissions_dir = Path(settings.BASE_DIR) / 'submission_json'
            submissions_dir.mkdir(exist_ok=True)
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            file_path = submissions_dir / f'deleted_submission_{submission.pk}_{timestamp}.json'
            # Serialize submission and answers
            data = SubmissionSerializer(submission).data
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            logger.info(f"[SubmissionViewSet.destroy] Wrote deleted submission to {file_path}")
        except Exception as file_exc:
            logger.error(f"[SubmissionViewSet.destroy] Failed to write deleted submission JSON: {file_exc}")
        # Delete the submission
        return super().destroy(request, *args, **kwargs)

    @action(detail=True, methods=['get'])
    def generate_doc(self, request, pk=None):
        submission = get_object_or_404(Submission, pk=pk)
        answers = {a.question.pk: a.value for a in Answer.objects.filter(submission=submission)}
        questions = Question.objects.filter(form=submission.form).order_by('order')
        form = submission.form
        doc = Document()

        # Add header with submission date and client name
        submission_date = submission.submission_date.strftime('%B %d, %Y') if submission.submission_date else ''
        client_name = submission.client_name or 'client'

        # Add header (client above date)
        header = doc.sections[0].header

        # Client name first
        p_client = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p_client.text = f"Client: {client_name}"
        p_client.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_client = p_client.runs[0] if p_client.runs else p_client.add_run()
        run_client.font.size = Pt(12)
        # Date second
        p_date = header.add_paragraph()
        p_date.text = f"Submission Date: {submission_date}"
        p_date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_date = p_date.runs[0] if p_date.runs else p_date.add_run()
        run_date.font.size = Pt(11)

        # Add "Personal circumstances" section
        doc.add_heading('Personal Circumstances', level=1)

        # Get form sections
        sections = form.sections or []

        # Initialize answer counter
        answer_number = 1

        # Iterate through sections
        for section in sections:
            # Add section heading
            doc.add_heading(section['name'], level=2)

            # Iterate through questions in this section
            for q in questions.filter(section_id=section['id']):
                template = q.output_template or ''
                answer = answers.get(q.pk, '')

                # Only add template if answer is not blank/null
                if answer is not None and str(answer).strip() != '':
                    text = f"{answer_number}. {template.replace('{{answer}}', str(answer))}"
                    if text.strip():
                        doc.add_paragraph(text)
                        answer_number += 1

        from io import BytesIO
        f = BytesIO()
        doc.save(f)
        f.seek(0)

        # Sanitize client name for filename
        safe_client_name = re.sub(r'[^a-zA-Z0-9_-]', '_', client_name)
        date_str = submission.submission_date.strftime('%Y%m%d') if submission.submission_date else ''
        filename = f"submission_{safe_client_name}_{date_str}_{submission.pk}.docx"
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Access-Control-Expose-Headers'] = 'Content-Disposition'
        return response

    def create(self, request, *args, **kwargs):
        logger = logging.getLogger(__name__)
        logger.info(f"[SubmissionViewSet.create] Incoming data: {json.dumps(request.data, indent=2)}")
        # Write the incoming submission data to a local JSON file for debugging
        try:
            from django.conf import settings
            from pathlib import Path
            submissions_dir = Path(settings.BASE_DIR) / 'submission_json'
            submissions_dir.mkdir(exist_ok=True)
            import datetime
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            file_path = submissions_dir / f'submission_{timestamp}.json'
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(request.data, f, ensure_ascii=False, indent=2)
            logger.info(f"[SubmissionViewSet.create] Wrote submission to {file_path}")
        except Exception as file_exc:
            logger.error(f"[SubmissionViewSet.create] Failed to write submission JSON: {file_exc}")
        try:
            response = super().create(request, *args, **kwargs)
            logger.info(f"[SubmissionViewSet.create] Response: {response.status_code} {getattr(response, 'data', None)}")
            return response
        except Exception as e:
            logger.error(f"[SubmissionViewSet.create] Exception: {str(e)}", exc_info=True)
            from rest_framework.views import exception_handler
            resp = exception_handler(e, context={'view': self, 'request': request})
            logger.error(f"[SubmissionViewSet.create] DRF exception handler response: {getattr(resp, 'data', None)}")
            return resp or Response({'detail': str(e)}, status=400)

    def update(self, request, *args, **kwargs):
        logger = logging.getLogger(__name__)
        logger.info(f"[SubmissionViewSet.update] Incoming data: {json.dumps(request.data, indent=2)}")
        # Write the incoming update data to a local JSON file for debugging
        try:
            submissions_dir = Path(settings.BASE_DIR) / 'submission_json'
            submissions_dir.mkdir(exist_ok=True)
            import datetime
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            file_path = submissions_dir / f'update_submission_{kwargs.get('pk', 'unknown')}_{timestamp}.json'
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(request.data, f, ensure_ascii=False, indent=2)
            logger.info(f"[SubmissionViewSet.update] Wrote update to {file_path}")
        except Exception as file_exc:
            logger.error(f"[SubmissionViewSet.update] Failed to write update JSON: {file_exc}")
        try:
            response = super().update(request, *args, **kwargs)
            logger.info(f"[SubmissionViewSet.update] Response: {response.status_code} {getattr(response, 'data', None)}")
            return response
        except Exception as e:
            logger.error(f"[SubmissionViewSet.update] Exception: {str(e)}", exc_info=True)
            from rest_framework.views import exception_handler
            resp = exception_handler(e, context={'view': self, 'request': request})
            logger.error(f"[SubmissionViewSet.update] DRF exception handler response: {getattr(resp, 'data', None)}")
            return resp or Response({'detail': str(e)}, status=400)