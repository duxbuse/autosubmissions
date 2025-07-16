from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from datetime import datetime

def add_title_page(doc, content, client_name):
    """Add a title page to the document following court format"""
    # Set up styles for left-aligned text
    style = doc.styles.add_style('Court Header Left', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    style.paragraph_format.space_after = Pt(0)

    # Add court name and location in top left
    p = doc.add_paragraph(style='Court Header Left')
    p.add_run(content['court'])
    
    p = doc.add_paragraph(style='Court Header Left')
    p.add_run(content['location'])

    # Set up center style for parties
    style = doc.styles.add_style('Court Header Center', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.bold = True
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    style.paragraph_format.space_after = Pt(6)

    # Add spacing before parties
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Center align the parties
    p = doc.add_paragraph(style='Court Header Center')
    p.add_run("VICTORIA POLICE")
    
    p = doc.add_paragraph(style='Court Header Center')
    p.add_run("and")

    # Add client name
    p = doc.add_paragraph(style='Court Header Center')
    p.add_run(client_name.upper())

    # Add spacing before title
    doc.add_paragraph()
    doc.add_paragraph()

    # Set up centered style for title
    if 'Court Header Center Bold' not in doc.styles:
        style = doc.styles.add_style('Court Header Center Bold', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_after = Pt(12)

    # Add the underlined title
    p = doc.add_paragraph(style='Court Header Center Bold')
    title_run = p.add_run("OUTLINE OF SUBMISSIONS FOR PLEA HEARING")
    title_run.bold = True
    title_run.underline = True
    
    # Add the information table
    doc.add_paragraph()  # Space before table
    table = doc.add_table(rows=4, cols=2)
    table.style = None  # Remove default table style
    table.autofit = False
    table.allow_autofit = False
    
    # Set full width table with right-aligned second column
    section = doc.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin
    # Convert to twips (twentieth of a point)
    total_width = int(available_width * 1440 / Inches(1))
    col1_width = int(total_width * 0.7)  # 70% for first column
    col2_width = int(total_width * 0.3)  # 30% for second column
    
    # Set column widths
    for cell in table.columns[0].cells:
        cell._tc.tcPr.append(parse_xml(f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{col1_width}" w:type="dxa"/>'))
    for cell in table.columns[1].cells:
        cell._tc.tcPr.append(parse_xml(f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{col2_width}" w:type="dxa"/>'))

    # Create table properties
    tblPr = parse_xml(r'''<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:tblBorders>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:right w:val="none"/>
            <w:insideH w:val="none"/>
            <w:insideV w:val="dashed" w:sz="4" w:space="0" w:color="auto"/>
        </w:tblBorders>
    </w:tblPr>''')
    
    # Remove any existing table properties
    for element in table._element.findall('.//w:tblPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        element.getparent().remove(element)
    
    # Add the new table properties
    table._element.insert(0, tblPr)

    # Fill table
    table.cell(0, 0).text = "Date of Document:"
    table.cell(0, 1).text = datetime.now().strftime("%d %B %Y")
    table.cell(1, 0).text = "Filed on behalf of:"
    table.cell(1, 1).text = "The Accused"
    table.cell(2, 0).text = "Prepared by:"
    table.cell(2, 1).text = "Solicitors code: 113 758"
    table.cell(3, 0).text = "James Dowsley & Associates"
    table.cell(3, 1).text = "Telephone: 9781 4900"

    # Format table text and align second column to right
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if i == 1:  # Second column
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

def add_section(doc, section_name, subsections_or_config, questions, questions_by_section, answers_by_question, sections):
    """Add a section with subsections and their fields"""
    # Create section heading style
    if 'Section Heading' not in doc.styles:
        style = doc.styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    # Create italic subheading style
    if 'Italic Subheading' not in doc.styles:
        style = doc.styles.add_style('Italic Subheading', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.font.italic = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Inches(0)

    # Create answer text style
    if 'Answer Text' not in doc.styles:
        style = doc.styles.add_style('Answer Text', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent for numbers
        style.paragraph_format.line_spacing = 1.0

    # Add the section heading
    p = doc.add_paragraph(style='Section Heading')
    p.add_run(section_name.upper())  # Make section heading uppercase
    
    # Initialize answer counter
    answer_number = 1
    
    # Process sections based on configuration
    if isinstance(subsections_or_config, dict) and subsections_or_config.get('use_form_sections'):
        # Use sections from the form
        for section in sections:
            section_name = section.get('name', '')
            if section_name:
                section_questions = [
                    q for q in questions.values()
                    if q.section_id == section.get('id') and answers_by_question.get(q.pk)
                ]
                
                if section_questions:
                    if section_name.lower() != "personal circumstances":
                        p = doc.add_paragraph(style='Italic Subheading')
                        p.add_run(section_name.lower())
                    
                    for q in section_questions:
                        answer_obj = answers_by_question.get(q.pk)
                        if answer_obj and str(answer_obj.value).strip():
                            p = doc.add_paragraph(style='Answer Text')
                            answer_text = f"{answer_number}. {q.text} {str(answer_obj.get_formatted_value())}"
                            p.add_run(answer_text)
                            answer_number += 1
