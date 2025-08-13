from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import os
from io import BytesIO
from datetime import datetime
from .models import Submission, Question, Answer

class DocGenerator:
    # Default template configurations for different form types
    DEFAULT_TEMPLATES = {
        'plea_of_guilty': {
            'title': "OUTLINE OF SUBMISSIONS FOR PLEA HEARING",
            'sections': [
                {'name': 'title_page', 'type': 'title', 'content': {'court': "IN THE MAGISTRATES' COURT", 'location': "AT MELBOURNE", 'parties': ["VICTORIA POLICE", "and", "{client_name}"]}},
                {'name': 'Chronology', 'type': 'section'},
                {'name': 'Material Tendered', 'type': 'section'},
                {'name': 'PRE-SENTENCE DETENTION CALCULATIONS', 'type': 'section'},
                {'name': 'ULTIMATE DISPOSITION', 'type': 'section'},
                {'name': 'CIRCUMSTANCES OF THE OFFENDING', 'type': 'section'},
                {'name': 'Personal Circumstances', 'type': 'section', 'use_form_sections': True},
                {'name': 'SENTENCING CONSIDERATIONS AND MATTERS IN MITIGATION', 'type': 'section'},
                {'name': 'DISPOSITION SOUGHT', 'type': 'section'},
                {'name': 'Signature', 'type': 'section', 'content': {'signature': "Ms Courtney Robertson\nLegal Representative for {{name}}", 'image': "signature.png", 'date': datetime.now().strftime('%Y-%m-%d')}},
            ]
        },
        'bail_application': {
            'title': "BAIL APPLICATION SUBMISSION",
            'sections': [
                {'name': 'title_page','type':'title','content': {'court': "IN THE MAGISTRATES' COURT", 'location': "AT MELBOURNE", 'parties': ["VICTORIA POLICE","and","{client_name}"]}},
                {'name': 'Application Details','type':'section','fields':['Charges','Exceptional Circumstances','Risk Assessment']},
            ]
        }
    }

    def __init__(self, submission: Submission):
        """Initialize with a submission"""
        self.submission = submission
        # Get all questions for this submission's form
        questions = Question.objects.filter(form=submission.form).order_by('order')
        self.questions = {q.pk: q for q in questions}
        self.sections = submission.form.sections or []

        # Get all questions by section for easy lookup
        self.questions_by_section = {}
        for q in questions:
            if q.section_id:
                if q.section_id not in self.questions_by_section:
                    self.questions_by_section[q.section_id] = []
                self.questions_by_section[q.section_id].append(q)

        # Get all answers for this submission
        answers = Answer.objects.filter(submission=submission)
        # Index answers by question for easy lookup
        self.answers_by_question = {a.question.pk: a for a in answers}
        for q in questions:
            if q.section_id:
                if q.section_id not in self.questions_by_section:
                    self.questions_by_section[q.section_id] = []
                self.questions_by_section[q.section_id].append(q)
        self.sections = submission.form.sections or []

    def generate(self) -> BytesIO:
        """Generate a document based on the form's template type"""
        template_type = self.submission.form.template_type or 'plea_of_guilty'
        template_config = self.DEFAULT_TEMPLATES.get(template_type, self.DEFAULT_TEMPLATES['plea_of_guilty'])
        
        doc = Document()
        
        # Set up the document
        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Set up default paragraph style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_after = Pt(0)  # No extra space after paragraphs by default
        
        # No header section needed

        # Get client name for the document (and remove honorifics)
        client_name = self.submission.client_name or 'client'
        print(f"Generating document for client: {client_name}")
        client_names = client_name.strip().split()
        client_name = ' '.join(client_names[-2:])  # Join names back together but only last 2 which should ignore the 1st honorific if present

        # Process template sections
        for section in template_config['sections']:
            if section['type'] == 'title':
                self._add_title_page(doc, section['content'], client_name)
            elif section['type'] == 'section':
                self._add_section(doc, section['name'], section)

        # Save to BytesIO
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def _add_title_page(self, doc, content, client_name):
        """Add a title page to the document based on the template type"""
        if self.submission.form.template_type == 'plea_of_guilty':
            from .templates.plea_of_guilty import add_title_page
            add_title_page(doc, content, client_name)
        else:
            # Default simple title page or other templates
            p = doc.add_paragraph()
            p.add_run(content.get('title', '')).bold = True

    def _add_section(self, doc, section_name, subsections_or_config):
        """Add a section with subsections and their fields based on the template type"""
        # Handle the signature section specifically
        if section_name == 'Signature':
            content = subsections_or_config.get('content', {})
            signature_text = content.get('signature','')
            signature_image = content.get('image','')
            date_str = datetime.now().strftime('%d/%m/%Y')

            client_name_for_sig = self.submission.client_name or 'client'
            signature_text = signature_text.replace('{{name}}', client_name_for_sig)

            p = doc.add_paragraph()
            p.add_run(signature_text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            if signature_image:
                base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
                image_path = os.path.join(base_dir, 'static', signature_image)
                try:
                    r = p.add_run()
                    r.add_picture(image_path, width=Inches(1.5))
                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                except Exception:
                    pass

            p_date = doc.add_paragraph()
            p_date.add_run(date_str)
            p_date.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            return

        if self.submission.form.template_type == 'plea_of_guilty':
            from .templates.plea_of_guilty import add_section
            add_section(doc, section_name, subsections_or_config, self.questions,
                        self.questions_by_section, self.answers_by_question, self.sections)
        else:
            answer_number = 1
            for subsection in subsections_or_config:
                subsection_lower = subsection.lower()
                matching_questions = [
                    q for q in self.questions.values()
                    if subsection_lower in q.text.lower()
                    and self.answers_by_question.get(q.pk)
                ]
                if matching_questions:
                    if subsection.lower() != "personal circumstances":
                        p = doc.add_paragraph(style='Italic Subheading')
                        p.add_run(subsection.lower())
                    else:
                        p = doc.add_paragraph(style='Section Heading')
                        p.add_run(subsection.upper())

                    for q in matching_questions:
                        answer_obj = self.answers_by_question.get(q.pk)
                        if answer_obj and str(answer_obj.value).strip():
                            p = doc.add_paragraph(style='Answer Text')
                            answer_text = f"{answer_number}. {q.text} {str(answer_obj.get_formatted_value())}"
                            p.add_run(answer_text)
                            answer_number += 1